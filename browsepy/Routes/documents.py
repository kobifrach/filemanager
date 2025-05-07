import os
import re
import pyodbc
import shutil
import uuid
from flask import Blueprint, request, jsonify
from pymongo import MongoClient
from docx import Document
from docx.shared import Pt
import win32com.client
from docx.enum.text import WD_COLOR_INDEX
from ..database import get_db_connection
from datetime import datetime
from docx.enum.text import WD_COLOR_INDEX
from flask import send_file
from browsepy.Utils.decorators import safe_route


# ---------- General Definitions ----------
documents_bp = Blueprint('documents', __name__)

# Connect to MongoDB
mongo_client = MongoClient("mongodb://localhost:27017")
mongo_db = mongo_client["Manage_Files_DB"]
values_collection = mongo_db["Customers_Fields"]

conn = get_db_connection()
cursor = conn.cursor()

VARIABLE_REGEX = r"\{\{(.*?)\}\}"

# ---------- Functions ----------

def validate_value_by_type(value, sql_type, field_name=None, errors=None):    


    #check if field is ID and if value is not number, return "_____"
    if 'id' in (field_name or '').lower():
        print("ID field found")
        if not str(value).isdigit():
            print("error in ID field")
            msg = f"שגיאה בשדה '{field_name or 'שדה'}': ערך '{value}' לא מספרי (שדה ID חייב להכיל רק ספרות)"
            if errors is not None:
                errors.append(msg)
            # if id not number, return _____ string
            return "_____"
        
    #check if field is match to its type in SQL
    type_match = re.match(r"([a-zA-Z]+)\(?(\d*)\)?", sql_type.strip())
    if not type_match:
        return value

    base_type = type_match.group(1).lower()
    length_str = type_match.group(2)
    max_length = int(length_str) if length_str.isdigit() else None

    try:
        if base_type in ['varchar', 'char', 'nchar', 'nvarchar']:
            str_value = str(value)
            if max_length and len(str_value) > max_length:
                return str_value[:max_length]
            return str_value

        elif base_type in ['int', 'bigint', 'smallint']:
            return int(value)

        elif base_type in ['float', 'real', 'decimal', 'numeric']:
            return float(value)

        elif base_type in ['bit']:
            return bool(int(value))

        elif base_type in ['date', 'datetime', 'smalldatetime']:
            return datetime.fromisoformat(str(value))

        else:
            return str(value)

    except Exception as e:
        msg = f"שגיאה בשדה '{field_name or 'שדה'}': ערך לא תקין '{value}' לסוג '{sql_type}'"
        if errors is not None:
            errors.append(msg)
            return "_____"
        else:
            return "_____"
            # raise ValueError(msg)


def convert_doc_to_docx(input_path):
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(input_path)
        output_path = os.path.splitext(input_path)[0] + ".docx"
        doc.SaveAs(output_path, FileFormat=16)  # 16 is docx format
        doc.Close()
        word.Quit()
        return output_path
    except Exception as e:
        print("שגיאה בהמרה ל-docx:", e)
        return None

#for downling
def convert_docx_to_pdf(docx_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(docx_path)

    temp_pdf_path = os.path.splitext(docx_path)[0] + "_temp.pdf"
    doc.SaveAs(temp_pdf_path, FileFormat=17)  # PDF
    doc.Close()
    word.Quit()

    return temp_pdf_path

def clean_text(text):
    """ניקוי טקסט מרווחים וסימני פיסוק כפולים"""
    text = re.sub(r'\s+', ' ', text)  # רווחים מיותרים
    text = re.sub(r'([.,;:!?])\s*\1+', r'\1', text)  # מחיקת כפילויות של סימני פיסוק
    text = re.sub(r'\s*([.,;:!?])', r'\1', text)  # מחיקת רווח לפני סימני פיסוק
    return text.strip()


def process_paragraphs(paragraphs, cursor, variables_not_found, sql_id, values):
    for para in paragraphs:
        matches = list(re.finditer(VARIABLE_REGEX, para.text))
        if not matches:
            continue

        original_text = para.text
        para.clear()
        last_end = 0

        for match in matches:
            start, end = match.span()
            var_text = match.group(1)
            parts = var_text.split(',')

            before = original_text[last_end:start]
            if before:
                before = re.sub(r'\s+', ' ', before)  # ניקוי רווחים מיותרים
                run = para.add_run(before)
                run.font.name = 'Arial'
                run.font.size = Pt(12)

            if len(parts) < 5:
                run = para.add_run(match.group(0))
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                last_end = end
                continue

            var_name, var_type, var_len, var_required, var_description = [p.strip() for p in parts]

            # Check if field exists in SQL, if not insert it
            cursor.execute("SELECT COUNT(*) FROM Fields WHERE name = ?", (var_name,))
            if cursor.fetchone()[0] == 0:
                cursor.execute("INSERT INTO Fields (name, type, length, required,description) VALUES (?, ?, ?, ?, ?)",
                               (var_name, var_type, int(var_len), int(var_required == '*'), var_description))
                cursor.connection.commit()

            # בדיקה האם השדה קיים ב-SQL
            cursor.execute("SELECT required FROM Fields WHERE name = ?", (var_name,))
            row = cursor.fetchone()
            field_is_required = row[0] == 1 if row else False

            value = values.get(var_name)
            field_errors = []

            if value is not None:
                sql_type = f"{var_type}({var_len})" if var_len else var_type
                validated_value = validate_value_by_type(value, sql_type, field_name=var_name, errors=field_errors)

                if validated_value is not None and validated_value != "_____":
                    run = para.add_run(str(validated_value))
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                else:
                    variables_not_found.append(var_name)
                    if field_is_required:
                        run = para.add_run("_____")
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                    else:
                        # לא להוסיף כלום ולא רווח
                        pass
            else:
                variables_not_found.append(var_name)
                if field_is_required:
                    run = para.add_run("_____")
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                else:
                    # לא להוסיף כלום ולא רווח
                    pass

            last_end = end

        after = original_text[last_end:]
        if after:
            after = re.sub(r'\s+', ' ', after)  # ניקוי רווחים מיותרים
            run = para.add_run(after)
            run.font.name = 'Arial'
            run.font.size = Pt(12)


def process_docx(path, sql_id):
    values = values_collection.find_one({"sql_id": sql_id}) or {}

    doc = Document(path)
    variables_not_found = []

    # Body paragraphs
    process_paragraphs(doc.paragraphs, cursor, variables_not_found, sql_id, values)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, cursor, variables_not_found, sql_id, values)

    # Headers and footers
    for section in doc.sections:
        process_paragraphs(section.header.paragraphs, cursor, variables_not_found, sql_id, values)
        process_paragraphs(section.footer.paragraphs, cursor, variables_not_found, sql_id, values)

    # Save document
    dir_path = os.path.dirname(path)
    filename_wo_ext = os.path.splitext(os.path.basename(path))[0]
    output_path = os.path.join(dir_path, f"{filename_wo_ext} טופל.docx")
    doc.save(output_path)

    print("✅ קובץ נשמר כ:", output_path)
    if variables_not_found:
        print("⚠️ משתנים שלא נמצאו:", variables_not_found)
    return variables_not_found

# ---------- Flask Routes ----------

@documents_bp.route('/process/<sql_id>', methods=['POST'])
@safe_route
def process_document_route(sql_id):
    data = request.get_json()
    file_path = data.get("file_path")
    if not file_path:
        return jsonify({"message": "יש לספק נתיב לקובץ"}), 400

    try:
        if file_path.endswith(".doc"):
            file_path = convert_doc_to_docx(file_path)
            if not file_path:
                return jsonify({"message": "שגיאה בהמרת קובץ ל-docx"}), 500

        missing_vars = process_docx(file_path, sql_id)
        return jsonify({
            "message": "עיבוד הושלם",
            "missing_variables": missing_vars
        }), 200
    except Exception as e:
        return jsonify({"message": f"שגיאה: {str(e)}"}), 500
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)
        # if cursor:
        #     cursor.close()
        # if conn:
        #     conn.close()
        # if mongo_client:
        #     mongo_client.close()

        # אם יש צורך למחוק את הקובץ
        # if os.path.exists(file_path):
        #     os.remove(file_path)

#export PDF file
@documents_bp.route('/export/pdf/<sql_id>', methods=['POST'])
@safe_route
def export_document_as_pdf(sql_id):
    data = request.get_json()
    file_path = data.get("file_path")
    if not file_path:
        return jsonify({"message": "יש לספק נתיב לקובץ"}), 400

    temp_docx_path = None
    temp_pdf_path = None

    try:
        # המרת doc ל-docx אם צריך
        if file_path.endswith(".doc"):
            file_path = convert_doc_to_docx(file_path)
            if not file_path:
                return jsonify({"message": "שגיאה בהמרת קובץ ל-docx"}), 500

        # עיבוד הקובץ ויצירת docx חדש
        process_docx(file_path, sql_id)

        # יצירת נתיב של הקובץ החדש (שנקרא "טופל")
        filename_wo_ext = os.path.splitext(os.path.basename(file_path))[0]
        temp_docx_path = os.path.join(os.path.dirname(file_path), f"{filename_wo_ext} טופל.docx")

        # המרה ל־PDF
        temp_pdf_path = convert_docx_to_pdf(temp_docx_path)

        # החזרת הקובץ להורדה
        return send_file(temp_pdf_path, as_attachment=True)

    except Exception as e:
        return jsonify({"message": f"שגיאה: {str(e)}"}), 500

    finally:
        # ניקוי קבצים זמניים
        for path in [file_path, temp_docx_path, temp_pdf_path]:
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                except Exception:
                    pass