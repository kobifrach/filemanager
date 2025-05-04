import logging
import os
import os.path
import json
import base64
from urllib.parse import quote
from flask import redirect, render_template, request, send_file, url_for, abort
import mimetypes

from flask import Response, send_from_directory, stream_with_context, make_response
from werkzeug.exceptions import NotFound
from werkzeug.utils import safe_join  # זה התיקון!

from .appconfig import Flask
from .manager import PluginManager
from .file import Node, secure_filename
from .exceptions import OutsideRemovableBase, OutsideDirectoryBase, InvalidFilenameError, InvalidPathError
from . import compat
from . import __meta__ as meta
from docx import Document
from io import BytesIO
from flask import render_template_string
from flask import render_template
from .database import get_db_connection
from docx import Document
from flask import render_template_string
# ייבוא הראוטים והבלופרינטים
from .Routes.Customers import customers_bp
from .Routes.Files import files_bp
from .Routes.Folders_Files import foldersFiles_bp
from .Routes.Folders import folders_bp
from .Routes.Customers_Folders import customer_folders_bp
from .Routes.User_Type import userTypes_bp
from .Routes.Users import users_bp
#from .Routes.additional_functions import addFunctions_bp
from .Routes.Customers_Files import customer_files_bp
from .Routes.Documents import documents_bp
from .Routes.PdfDocuments import pdf_documents_bp
# from .Routes.ocr import ocr_bp



__app__ = meta.app  # noqa
__version__ = meta.version  # noqa
__license__ = meta.license  # noqa
__author__ = meta.author  # noqa
__basedir__ = os.path.abspath(os.path.dirname(compat.fsdecode(__file__)))

logger = logging.getLogger(__name__)

app = Flask(
    __name__,
    static_url_path='/static',
    static_folder=os.path.join(__basedir__, "static"),
    template_folder=os.path.join(__basedir__, "templates")
    )

app.register_blueprint(customers_bp)
app.register_blueprint(files_bp)
app.register_blueprint(foldersFiles_bp)
app.register_blueprint(folders_bp)
app.register_blueprint(customer_folders_bp)
app.register_blueprint(userTypes_bp)
app.register_blueprint(users_bp)
app.register_blueprint(documents_bp)
app.register_blueprint(pdf_documents_bp)
app.register_blueprint(customer_files_bp)
# app.register_blueprint(ocr_bp)

app.config.update(
    # directory_base=compat.getcwd(),
    directory_base=os.path.abspath('./files'),  # שמירה בתיקיית files
    # directory_start=None,
    directory_start=os.path.abspath('./files/start_folder'),  # תיקיית התחלה
    # directory_remove=None,
    directory_remove=os.path.abspath('./files/removable'),  # תיקיית מחיקה
    # directory_upload=None,
    directory_upload=os.path.abspath('./files/uploads'),  # תיקיית העלאה
    
    directory_tar_buffsize=262144,
    directory_downloadable=True,
    use_binary_multiples=True,
    plugin_modules=[],
    plugin_namespaces=(
        'browsepy.plugin',
        'browsepy_',
        '',
        ),
    exclude_fnc=None,
    )
app.jinja_env.add_extension('browsepy.transform.htmlcompress.HTMLCompress')

if 'BROWSEPY_SETTINGS' in os.environ:
    app.config.from_envvar('BROWSEPY_SETTINGS')

plugin_manager = PluginManager(app)

def iter_cookie_browse_sorting(cookies):
    '''
    Get sorting-cookie from cookies dictionary.

    :yields: tuple of path and sorting property
    :ytype: 2-tuple of strings
    '''
    try:
        data = cookies.get('browse-sorting', 'e30=').encode('ascii')
        for path, prop in json.loads(base64.b64decode(data).decode('utf-8')):
            yield path, prop
    except (ValueError, TypeError, KeyError) as e:
        logger.exception(e)



def get_cookie_browse_sorting(path, default):
    '''
    Get sorting-cookie data for path of current request.

    :returns: sorting property
    :rtype: string
    '''
    if request:
        for cpath, cprop in iter_cookie_browse_sorting(request.cookies):
            if path == cpath:
                return cprop
    return default

@app.route('/users')
def users():
    # חיבור למסד הנתונים
    conn = get_db_connection()
    cursor = conn.cursor()

    # שליפת נתונים מהטבלה
    cursor.execute('SELECT * FROM Users')
    rows = cursor.fetchall()

    # יצירת רשימה עם התוצאות
    users_list = []
    for row in rows:
        users_list.append({'id': row[0], 'name': row[1], 'email': row[2]})

    conn.close()

    return render_template('users.html', users=users_list)

@app.route('/sort/<string:property>', defaults={"path": ""})

@app.route('/sort/<string:property>/<path:path>')
def sort(property, path):
    try:
        directory = Node.from_urlpath(path)
    except OutsideDirectoryBase:
        return NotFound()

    if not directory.is_directory or directory.is_excluded:
        return NotFound()

    data = [
        (cpath, cprop)
        for cpath, cprop in iter_cookie_browse_sorting(request.cookies)
        if cpath != path
        ]
    data.append((path, property))
    raw_data = base64.b64encode(json.dumps(data).encode('utf-8')).decode('utf-8')  # שינוי כאן

    # prevent cookie becoming too large
    while len(raw_data) > 3975:  # 4000 - len('browse-sorting=""; Path=/')
        data.pop(0)
        raw_data = base64.b64encode(json.dumps(data).encode('utf-8')).decode('utf-8')  # שינוי כאן

    response = redirect(url_for(".browse", path=directory.urlpath))
    response.set_cookie('browse-sorting', raw_data)
    return response


@app.route("/browse", defaults={"path": ""})
@app.route('/browse/<path:path>')
def browse(path):
    sort_property = get_cookie_browse_sorting(path, 'text')
    sort_fnc, sort_reverse = browse_sortkey_reverse(sort_property)
    try:
        directory = Node.from_urlpath(path)
        if directory.is_directory and not directory.is_excluded:
            return stream_template(
                'browse.html',
                file=directory,
                sort_property=sort_property,
                sort_fnc=sort_fnc,
                sort_reverse=sort_reverse
                )
    except OutsideDirectoryBase:
        pass
    return NotFound()


@app.route("/preview/file/<path:path>")
def preview_file(path):
    try:
        file = Node.from_urlpath(path)

        if not file.is_file or file.is_excluded:
            return NotFound()

        file_path = getattr(file, "path", None) or getattr(file, "fullpath", None) or getattr(file, "absolute_path", None)

        if not file_path:
            return NotFound()

        # אם הקובץ הוא DOCX, נמיר אותו ל-HTML
        if file_path.endswith(".docx"):
            doc = Document(file_path)
            html_content = "<html><body>"
            for para in doc.paragraphs:
                html_content += f"<p>{para.text}</p>"
            html_content += "</body></html>"
            return render_template_string(html_content)

        # אחרת, נשלח אותו כרגיל עם MIME type מתאים
        mimetype, _ = mimetypes.guess_type(file_path)
        return send_file(file_path, mimetype=mimetype, as_attachment=False)

    except Exception as e:
        print(f"Error: {e}")
        return NotFound()


@app.route('/open/<path:path>', endpoint="open")
def open_file(path):
    try:
        file = Node.from_urlpath(path)
        if file.is_file and not file.is_excluded:
            return send_from_directory(file.parent.path, file.name)
    except OutsideDirectoryBase:
        pass
    return NotFound()


@app.route("/download/file/<path:path>")
def download_file(path):
    try:
        file = Node.from_urlpath(path)
        if file.is_file and not file.is_excluded:
            return file.download()
    except OutsideDirectoryBase:
        pass
    return NotFound()


@app.route("/download/directory/<path:path>.tgz")
def download_directory(path):
    try:
        directory = Node.from_urlpath(path)
        if directory.is_directory and not directory.is_excluded:
            return directory.download()
    except OutsideDirectoryBase:
        pass
    return NotFound()


@app.route("/create_folder", methods=["POST"])
def create_folder():
    UPLOAD_FOLDER = r"C:\Users\User\Desktop\files"
    try:
        folder_name = request.form.get('folder_name')
        current_path = request.form.get('current_path')  # קבלת הנתיב מהבקשה

        if not folder_name:
            return "Folder name is required", 400

        if not current_path:
            current_path = "uploads"  # ברירת מחדל אם אין נתיב
        
        folder_path = os.path.join(UPLOAD_FOLDER, current_path, folder_name)

        os.makedirs(folder_path, exist_ok=True)

        doc = Document()
        doc.add_heading(folder_name, level=1)
        doc.save(os.path.join(folder_path, f"{folder_name}.docx"))

        return redirect(url_for(".browse", path=current_path))

    except Exception as e:
        app.logger.error(f"Error during folder creation: {str(e)}")
        return "An error occurred", 500




@app.route("/remove/<path:path>", methods=("GET", "POST"))
def remove(path):
    try:
        file = Node.from_urlpath(path)
        file.can_remove = True

    except OutsideDirectoryBase:
        return NotFound()

    if not file.can_remove or file.is_excluded or not file.parent:
        return NotFound()

    if request.method == 'GET':
        return render_template('remove.html', file=file)

    file.remove()
    return redirect(url_for(".browse", path=file.parent.urlpath))



@app.route("/upload", defaults={'path': ''}, methods=("POST",))
@app.route("/upload/<path:path>", methods=("POST",))
def upload(path):
    UPLOAD_FOLDER = r"C:\Users\User\Desktop\files"
    try:
        print(f"Received path: {path}")  # לוג עם הנתיב המתקבל
        
        if not path:  # אם הנתיב ריק, נגדיר ברירת מחדל
            path = "uploads"  # ברירת מחדל אם אין נתיב
        print(f"Using path (default if needed): {path}")  # לוג אחרי קביעת ברירת המחדל

        # יצירת נתיב מלא
        folder_path = os.path.join(UPLOAD_FOLDER, path)
        print(f"Folder path: {folder_path}")  # לוג של הנתיב המלא

        # אם התיקייה לא קיימת, צור אותה
        if not os.path.exists(folder_path):
            app.logger.info(f"Creating folder: {folder_path}")
            os.makedirs(folder_path)
        
        directory = Node.from_urlpath(path)  # השתמש ב-Node כפי שעשית קודם
        directory.can_upload = True
        print(f"Directory: {directory}")  # לוג של אובייקט הדירקטוריה

        # הדפסת ערכים שמאפשרים לך להבין אם התנאים נכונים
        print(f"Is directory: {directory.is_directory}")
        print(f"Can upload: {directory.can_upload}")
        print(f"Is excluded: {directory.is_excluded}")

        if not directory.is_directory or not directory.can_upload or directory.is_excluded:
            print("Invalid directory")
            return NotFound()
        
        print(f"Files in request: {request.files}")

        for f in request.files.getlist('file[]'):
            filename = secure_filename(f.filename)
            
            if filename:
                print(f"filename: {filename}")  # לוג של שם הקובץ
                filename = directory.choose_filename(filename)
                filepath = os.path.join(folder_path, filename)  # ודא שאתה משתמש ב-folder_path
                print(f"Saving file to: {filepath}")  # לוג שמראה היכן נשמר הקובץ
                f.save(filepath)
            else:
                print(f"Invalid filename: {f.filename}")
                raise InvalidFilenameError(
                    path=directory.path,
                    filename=f.filename
                )
        print("File uploaded successfully")

        return redirect(url_for(".browse", path=directory.urlpath))
    
    except Exception as e:
        app.logger.error(f"Error during file upload: {str(e)}")  # לוג של שגיאה במקרה של בעיה
        print(f"Error during file upload: {str(e)}")  # לוג של שגיאה במקרה של בעיה
        return "An error occurred", 500



def browse_sortkey_reverse(prop):
    '''
    Get sorting function for directory listing based on given attribute
    name, with some caveats:
    * Directories will be first.
    * If *name* is given, link widget lowercase text will be used istead.
    * If *size* is given, bytesize will be used.

    :param prop: file attribute name
    :returns: tuple with sorting gunction and reverse bool
    :rtype: tuple of a dict and a bool
    '''
    if prop.startswith('-'):
        prop = prop[1:]
        reverse = True
    else:
        reverse = False

    if prop == 'text':
        return (
            lambda x: (
                x.is_directory == reverse,
                x.link.text.lower() if x.link and x.link.text else x.name
                ),
            reverse
            )
    if prop == 'size':
        return (
            lambda x: (
                x.is_directory == reverse,
                x.stats.st_size
                ),
            reverse
            )
    return (
        lambda x: (
            x.is_directory == reverse,
            getattr(x, prop, None)
            ),
        reverse
        )


def stream_template(template_name, **context):
    '''
    Some templates can be huge, this function returns an streaming response,
    sending the content in chunks and preventing from timeout.

    :param template_name: template
    :param **context: parameters for templates.
    :yields: HTML strings
    '''
    app.update_template_context(context)
    template = app.jinja_env.get_template(template_name)
    stream = template.generate(context)
    return Response(stream_with_context(stream))


@app.context_processor
def template_globals():
    return {
        'manager': app.extensions['plugin_manager'],
        'len': len,
        }


@app.route("/")
def index():
    path = app.config["directory_start"] or app.config["directory_base"]

    # הדפסות לבדיקה
    print("Config directory_start:", app.config.get("directory_start"))
    print("Config directory_base:", app.config.get("directory_base"))
    print("Final path used:", path)

    try:
        urlpath = Node(path).urlpath
    except OutsideDirectoryBase:
        print("Error: Path is outside the directory base!")
        return NotFound()

    return browse(urlpath)


@app.after_request
def page_not_found(response):
    if response.status_code == 404:
        return make_response((render_template('404.html'), 404))
    return response


@app.errorhandler(InvalidPathError)
def bad_request_error(e):
    file = None
    if hasattr(e, 'path'):
        if isinstance(e, InvalidFilenameError):
            file = Node(e.path)
        else:
            file = Node(e.path).parent
    return render_template('400.html', file=file, error=e), 400


@app.errorhandler(OutsideRemovableBase)
@app.errorhandler(404)
def page_not_found_error(e):
    return render_template('404.html'), 404


@app.errorhandler(500)
def internal_server_error(e):  # pragma: no cover
    logger.exception(e)
    return getattr(e, 'message', 'Internal server error'), 500
