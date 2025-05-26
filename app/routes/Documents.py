import os
import re
import pyodbc
import shutil
import uuid
from flask import Blueprint, request, jsonify, current_app
from pymongo import MongoClient
from docx import Document
from docx.shared import Pt
import win32com.client
from docx.enum.text import WD_COLOR_INDEX
from ..database.database import get_db_connection
from datetime import datetime
from flask import send_file
from ..utils.decorators import safe_route
from ..utils.jwt_decorator import token_required


# ---------- General Definitions ----------
documents_bp = Blueprint('documents', __name__)

# Connect to MongoDB
mongo_client = MongoClient("mongodb://localhost:27017")
mongo_db = mongo_client["Manage_Files_DB"]
values_collection = mongo_db["Customers_Fields"]

conn = get_db_connection()
cursor = conn.cursor()

VARIABLE_REGEX = r"\{\{(.*?)\}\}"

# ---------- Functions ----------

def validate_value_by_type(value, sql_type, field_name=None, errors=None):
    # Validates the value according to its SQL type
    
    # Check if field is ID and if value is not a number, return "_____"
    if 'id' in (field_name or '').lower():
        current_app.logger.info(f"Checking ID field for value: {value}")
        if not str(value).isdigit():
            msg = f"Error in field '{field_name or 'Field'}': Value '{value}' is not numeric (ID fields must contain only numbers)"
            if errors is not None:
                errors.append(msg)
            return "_____"
        
    # Check if field matches its type in SQL
    type_match = re.match(r"([a-zA-Z]+)\(?(\d*)\)?", sql_type.strip())
    if not type_match:
        return value

    base_type = type_match.group(1).lower()
    length_str = type_match.group(2)
    max_length = int(length_str) if length_str.isdigit() else None

    try:
        if base_type in ['varchar', 'char', 'nchar', 'nvarchar']:
            str_value = str(value)
            if max_length and len(str_value) > max_length:
                return str_value[:max_length]
            return str_value

        elif base_type in ['int', 'bigint', 'smallint']:
            return int(value)

        elif base_type in ['float', 'real', 'decimal', 'numeric']:
            return float(value)

        elif base_type in ['bit']:
            return bool(int(value))

        elif base_type in ['date', 'datetime', 'smalldatetime']:
            return datetime.fromisoformat(str(value))

        else:
            return str(value)

    except Exception as e:
        msg = f"Error in field '{field_name or 'Field'}': Invalid value '{value}' for type '{sql_type}'"
        if errors is not None:
            errors.append(msg)
        return "_____"

def convert_doc_to_docx(input_path):
    # Converts DOC file to DOCX format
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(input_path)
        output_path = os.path.splitext(input_path)[0] + ".docx"
        doc.SaveAs(output_path, FileFormat=16)  # 16 is docx format
        doc.Close()
        word.Quit()
        current_app.logger.info(f"File converted to DOCX: {output_path}")
        return output_path
    except Exception as e:
        current_app.logger.error(f"Error converting to DOCX: {e}")
        return None

def convert_docx_to_pdf(docx_path):
    # Converts DOCX file to PDF format
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)

        temp_pdf_path = os.path.splitext(docx_path)[0] + "_temp.pdf"
        doc.SaveAs(temp_pdf_path, FileFormat=17)  # PDF
        doc.Close()
        word.Quit()
        current_app.logger.info(f"File converted to PDF: {temp_pdf_path}")
        return temp_pdf_path
    except Exception as e:
        current_app.logger.error(f"Error converting DOCX to PDF: {e}")
        return None

def clean_text(text):
    # Cleans the text from excessive spaces and repeated punctuation
    text = re.sub(r'\s+', ' ', text)  # Removing extra spaces
    text = re.sub(r'([.,;:!?])\s*\1+', r'\1', text)  # Removing repeated punctuation
    text = re.sub(r'\s*([.,;:!?])', r'\1', text)  # Removing space before punctuation
    return text.strip()

def process_paragraphs(paragraphs, cursor, variables_not_found, sql_id, values):
    # Processes paragraphs and replaces variables
    for para in paragraphs:
        matches = list(re.finditer(VARIABLE_REGEX, para.text))
        if not matches:
            continue

        original_text = para.text
        para.clear()
        last_end = 0

        for match in matches:
            start, end = match.span()
            var_text = match.group(1)
            parts = var_text.split(',')

            before = original_text[last_end:start]
            if before:
                before = re.sub(r'\s+', ' ', before)  # Cleaning extra spaces
                run = para.add_run(before)
                run.font.name = 'Arial'
                run.font.size = Pt(12)

            if len(parts) < 5:
                run = para.add_run(match.group(0))
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                last_end = end
                continue

            var_name, var_type, var_len, var_required, var_description = [p.strip() for p in parts]

            # Check if field exists in SQL, if not insert it
            cursor.execute("SELECT COUNT(*) FROM Fields WHERE name = ?", (var_name,))
            if cursor.fetchone()[0] == 0:
                cursor.execute("INSERT INTO Fields (name, type, length, required, description) VALUES (?, ?, ?, ?, ?)",
                               (var_name, var_type, int(var_len), int(var_required == '*'), var_description))
                cursor.connection.commit()

            # Check if the field is required
            cursor.execute("SELECT required FROM Fields WHERE name = ?", (var_name,))
            row = cursor.fetchone()
            field_is_required = row[0] == 1 if row else False

            value = values.get(var_name)
            field_errors = []

            if value is not None:
                sql_type = f"{var_type}({var_len})" if var_len else var_type
                validated_value = validate_value_by_type(value, sql_type, field_name=var_name, errors=field_errors)

                if validated_value is not None and validated_value != "_____":
                    run = para.add_run(str(validated_value))
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                else:
                    variables_not_found.append(var_name)
                    if field_is_required:
                        run = para.add_run("_____")
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                    else:
                        # Do not add anything or space
                        pass
            else:
                variables_not_found.append(var_name)
                if field_is_required:
                    run = para.add_run("_____")
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                else:
                    # Do not add anything or space
                    pass

            last_end = end

        after = original_text[last_end:]
        if after:
            after = re.sub(r'\s+', ' ', after)  # Cleaning extra spaces
            run = para.add_run(after)
            run.font.name = 'Arial'
            run.font.size = Pt(12)

def process_docx(path, sql_id):
    # Processes DOCX file and replaces variables with values from DB
    values = values_collection.find_one({"sql_id": sql_id}) or {}
    doc = Document(path)
    variables_not_found = []

    # Body paragraphs
    process_paragraphs(doc.paragraphs, cursor, variables_not_found, sql_id, values)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, cursor, variables_not_found, sql_id, values)

    # Headers and footers
    for section in doc.sections:
        process_paragraphs(section.header.paragraphs, cursor, variables_not_found, sql_id, values)
        process_paragraphs(section.footer.paragraphs, cursor, variables_not_found, sql_id, values)

    # Save the document
    dir_path = os.path.dirname(path)
    filename_wo_ext = os.path.splitext(os.path.basename(path))[0]
    output_path = os.path.join(dir_path, f"{filename_wo_ext} טופל.docx")
    doc.save(output_path)

    current_app.logger.info(f"Processed document saved as: {output_path}")
    if variables_not_found:
        current_app.logger.warning(f"Variables not found: {variables_not_found}")
    return variables_not_found

# ---------- Flask Routes ----------

@documents_bp.route('/process/<sql_id>', methods=['POST'])
@safe_route
@token_required()
def process_document_route(sql_id):
    # Route for processing DOCX file
    data = request.get_json()
    file_path = data.get("file_path")
    if not file_path:
        return jsonify({"message": "יש לספק נתיב לקובץ"}), 400

    try:
        if file_path.endswith(".doc"):
            file_path = convert_doc_to_docx(file_path)
            if not file_path:
                return jsonify({"message": "שגיאה בהמרת קובץ ל-docx"}), 500

        missing_vars = process_docx(file_path, sql_id)
        return jsonify({
            "message": "עיבוד הושלם",
            "missing_variables": missing_vars
        }), 200
    except Exception as e:
        current_app.logger.error(f"Error processing document: {e}")
        return jsonify({"message": f"שגיאה: {str(e)}"}), 500
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)


@documents_bp.route('/extract_variables', methods=['POST'])
@safe_route
@token_required()
def extract_variables_route():
    data = request.get_json()
    file_path = data.get("file_path")
    if not file_path:
        return jsonify({"message": "יש לספק נתיב לקובץ"}), 400

    try:
        if file_path.endswith(".doc"):
            file_path = convert_doc_to_docx(file_path)
            if not file_path:
                return jsonify({"message": "שגיאה בהמרת קובץ ל-docx"}), 500

        doc = Document(file_path)  # טען את הקובץ כאן
        variables_found = set()

        # Process paragraphs to extract variables
        for para in doc.paragraphs:
            matches = re.findall(VARIABLE_REGEX, para.text)
            for match in matches:
                parts = match.split(',')
                if parts:
                    variables_found.add(parts[0].strip())

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = re.findall(VARIABLE_REGEX, para.text)
                        for match in matches:
                            parts = match.split(',')
                            if parts:
                                variables_found.add(parts[0].strip())

        # Convert set to list
        variables_list = list(variables_found)

        return jsonify({
            "message": "משתנים נמשכו בהצלחה",
            "variables": variables_list
        }), 200
    except Exception as e:
        current_app.logger.error(f"Error extracting variables: {e}")
        return jsonify({"message": f"שגיאה: {str(e)}"}), 500

# @documents_bp.route('/export/pdf/<sql_id>', methods=['POST'])
# @safe_route
# def export_document_as_pdf(sql_id):
#     # Route for exporting document as PDF
#     data = request.get_json()
#     file_path = data.get("file_path")
#     if not file_path:
#         return jsonify({"message": "יש לספק נתיב לקובץ"}), 400

#     temp_docx_path = file_path + ".docx"
#     missing_vars = process_docx(file_path, sql_id)

#     # If no missing variables, proceed with conversion to PDF
#     if not missing_vars:
#         pdf_path = convert_docx_to_pdf(temp_docx_path)
#         if pdf_path:
#             return send_file(pdf_path, as_attachment=True)
#         else:
#             return jsonify({"message": "שגיאה בהמרת קובץ ל-PDF"}), 500
#     else:
#         return jsonify({"message": "חסרים משתנים, לא ניתן להפיק PDF"}), 400
# @documents_bp.route('/download', methods=['POST'])
# @safe_route
# def download_file():
#     # Route for downloading a file given its path
#     data = request.get_json()
#     file_path = data.get("file_path")  # Get the file path from the request body
    
#     if not file_path:
#         return jsonify({"message": "יש לספק נתיב לקובץ"}), 400
    
#     # Ensure the file exists at the given path
#     if not os.path.isfile(file_path):
#         return jsonify({"message": "הקובץ לא נמצא בנתיב שסופק"}), 404
    
#     # Send the file to the client as an attachment
#     return send_file(file_path, as_attachment=True)

@documents_bp.route('/download', methods=['POST'])
@token_required
def download_pdf():
    # Assuming you pass only the filename (without path)
    data = request.get_json()
    file_name = data.get("file_name")
    
    if not file_name:
        return jsonify({"message": "יש לספק שם קובץ"}), 400

    # Assuming file is located in the same directory as the Flask app
    file_path = os.path.join(os.getcwd(), file_name)

    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return jsonify({"message": "הקובץ לא נמצא"}), 404